#!/usr/bin/env python3
"""Create the implementation-baseline MascotRender SDD v0.2 from v0.1."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "MascotRender_SDD_v0.1.docx"
TARGET = ROOT / "MascotRender_SDD_v0.2.docx"


def paragraph(document: Document, exact_text: str):
    matches = [item for item in document.paragraphs if item.text == exact_text]
    if len(matches) > 1:
        heading_matches = [
            item for item in matches if item.style.name.startswith("Heading")
        ]
        if len(heading_matches) == 1:
            return heading_matches[0]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph matching {exact_text!r}; found {len(matches)}"
        )
    return matches[0]


def replace_paragraph(document: Document, old: str, new: str) -> None:
    paragraph(document, old).text = new


def set_cell(cell, text: str) -> None:
    cell.text = text


def set_code_cell(cell, text: str) -> None:
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8)


def set_row(table, row_index: int, values: tuple[str, ...]) -> None:
    row = table.rows[row_index]
    if len(row.cells) != len(values):
        raise RuntimeError(f"Table row has {len(row.cells)} cells, got {len(values)}")
    for cell, value in zip(row.cells, values, strict=True):
        set_cell(cell, value)


def update_requirement(table, requirement_id: str, priority: str, text: str) -> None:
    for row in table.rows[1:]:
        if row.cells[0].text == requirement_id:
            set_cell(row.cells[1], priority)
            set_cell(row.cells[2], text)
            return
    raise RuntimeError(f"Requirement {requirement_id} not found")


def append_requirement(table, requirement_id: str, priority: str, text: str) -> None:
    row = table.add_row()
    for cell, value in zip(
        row.cells, (requirement_id, priority, text), strict=True
    ):
        set_cell(cell, value)


def insert_before(document: Document, anchor_text: str, entries) -> None:
    anchor = paragraph(document, anchor_text)
    for text, style in entries:
        anchor.insert_paragraph_before(text, style)


def main() -> None:
    document = Document(SOURCE)

    # Cover, control, and revision metadata.
    replace_paragraph(
        document,
        "Document status: Draft for technical and product review",
        "Document status: Implementation baseline",
    )
    replace_paragraph(document, "Version: 0.1", "Version: 0.2")
    replace_paragraph(document, "Core design decision", "Core implementation decisions")
    document.tables[1].cell(0, 0).text = (
        "Local-first procedural generation; installable C++20 library\n"
        "MascotRender 0.1 will be a build-time content compiler distributed with "
        "Conan 2 and CMake. It will use a private ThorVG software backend and "
        "libwebp, accept versioned JSON specifications and trusted curated SVG "
        "packs, and generate deterministic static assets and manifests. Runtime "
        "chat suggestions consume generated data and remain outside the engine MVP."
    )

    control = document.tables[2]
    set_cell(control.cell(3, 1), "0.2")
    set_cell(control.cell(4, 1), "Implementation baseline")
    set_cell(control.cell(6, 1), "C++ Graphics/Platform Lead")
    set_cell(
        control.cell(8, 1),
        "Engine 0.1.0 after M0-M5; 50-sticker coherence gate and pilot follow",
    )

    revision = document.tables[3].add_row()
    for cell, value in zip(
        revision.cells,
        (
            "0.2",
            "13 Jul 2026",
            "C++ Technical Lead / Project Management",
            "Implementation baseline: separated engine MVP from product pilot; "
            "adopted Conan 2, ThorVG, JSON-first authoring, an installable public "
            "API, profile-scoped determinism, and milestone exit gates.",
        ),
        strict=True,
    ):
        set_cell(cell, value)

    document.tables[4].cell(0, 0).text = (
        "Recommended delivery boundary\n"
        "M0-M5 deliver an installable build-time C++20 engine. M6 validates a "
        "50-sticker art system. M7 delivers the approximately 200-sticker product "
        "pilot and runtime integration. CSV, full multilingual shaping, incremental "
        "builds, OpenCV QA, animation, Drogon, and React do not block engine 0.1.0."
    )

    # Executive summary and scope.
    replace_paragraph(
        document,
        "The engine is deliberately procedural rather than open-ended. A technical artist defines one or more mascot packs containing vector body parts, anchors, expressions, poses, accessories, palettes, text frames, and effects. Content authors then describe each sticker in CSV or JSON. MascotRender compiles those inputs into exact, reproducible WebP/SVG assets, thumbnails, a catalogue, a trigger dictionary, and a quality report.",
        "The engine is deliberately procedural rather than open-ended. For engine "
        "0.1, a technical artist defines a trusted versioned JSON mascot pack "
        "containing curated SVG parts, anchors, expressions, poses, effects, "
        "palettes, text frames, and a licensed font. Content is authored as JSON. "
        "MascotRender compiles these inputs into transparent WebP assets, thumbnails, "
        "a catalogue, a trigger dictionary, and a build report.",
    )
    success_updates = {
        "A content operator can generate a versioned pack of at least 1,000 static stickers from structured input on a normal developer workstation.":
            "An unrelated C++20 application can consume mascotrender/0.1.0 through Conan 2 and link MascotRender::MascotRender.",
        "The same sticker ID and pack version produce byte-identical or visually deterministic output across supported build environments.":
            "A test consumer loads an example pack and produces transparent 512 x 512 and 256 x 256 WebP assets.",
        "Displayed wording is rendered by the text engine exactly as authored; no model can misspell it.":
            "The same source, font files, lockfile, renderer settings, and pinned Conan profiles produce byte-identical output.",
        "The React client resolves suggestions locally with no per-keystroke network dependency.":
            "Ten representative stickers render exact English/Pidgin wording or fail with structured diagnostics.",
        "The asset library uses an original Wahalao visual system and avoids third-party characters, celebrity likenesses, or copied sticker artwork.":
            "Public headers expose no renderer, codec, JSON, or CLI dependency types; the installed package passes a separate test_package consumer.",
    }
    for old, new in success_updates.items():
        replace_paragraph(document, old, new)

    replace_paragraph(document, "3.3 MVP Scope", "3.3 Engine MVP Scope")
    scope = document.tables[6]
    scope_rows = [
        ("C++20 library, build/validate CLI, and Conan 2 package", "Full graphical authoring studio"),
        ("One original trusted layered SVG example pack", "Untrusted third-party pack ingestion and marketplace"),
        ("JSON sticker/pack input; WebP output and thumbnails", "CSV authoring workflow and binary Trie/FST"),
        ("Expressions, poses, effects, palettes, deterministic seed", "3D rendering, physics, and animation"),
        ("Exact English/Pidgin text with packaged font and fitting", "Full multilingual shaping, fallback, bidi, and Korean launch support"),
        ("Catalogue, dictionary, build report, validation", "Incremental cache, OpenCV similarity, and contact-sheet tooling"),
        ("Installable API plus external consumer test", "Drogon delivery, React tray, MLS integration, and product telemetry"),
    ]
    for index, values in enumerate(scope_rows, start=1):
        set_row(scope, index, values)

    # Requirements: MUST is engine 0.1 unless explicitly marked PILOT MUST.
    document.paragraphs[84].text = (
        "Unless a row is marked PILOT MUST, MUST refers to engine 0.1.0 (M0-M5). "
        "PILOT MUST applies to M6-M7 and does not block the installable engine release."
    )
    functional = document.tables[8]
    # In the source document the placeholder paragraph follows the table in XML.
    # Move the explanatory paragraph ahead of it so the scope qualifier is read first.
    functional._tbl.addprevious(document.paragraphs[84]._p)
    update_requirement(functional, "FR-01", "MUST", "Ingest JSON sticker specifications with schema-version validation and human-readable, source-located errors.")
    update_requirement(functional, "FR-08", "MUST", "Export transparent 512 x 512 WebP assets and 256 x 256 thumbnails; SVG source export is optional after engine 0.1.")
    update_requirement(functional, "FR-10", "SHOULD", "Perform dependency-aware incremental builds during the 50-sticker coherence milestone.")
    update_requirement(functional, "FR-11", "MUST", "Validate missing references, path containment, dimensions, alpha, visible bounds, text safe area, and file-size warnings. Perceptual similarity is deferred.")
    update_requirement(functional, "FR-12", "MUST", "Expose build and validate CLI commands. Preview, inspect, and package subcommands are later workflow enhancements.")
    update_requirement(functional, "FR-15", "PILOT MUST", "Allow React to resolve suggestions locally and return ordered sticker IDs without a per-keystroke request.")
    update_requirement(functional, "FR-16", "PILOT MUST", "Package runtime assets under an immutable, versioned directory suitable for CDN/Drogon delivery.")
    update_requirement(functional, "FR-18", "SHOULD", "Add HarfBuzz/FreeType shaping and locale-aware font fallback after the base English/Pidgin text pipeline is proven.")
    update_requirement(functional, "FR-19", "SHOULD", "Produce a static HTML or contact-sheet review page for the 50-sticker coherence gate.")
    append_requirement(functional, "FR-22", "SHOULD", "Add CSV import by converting operator-friendly rows into the canonical JSON/StickerSpec model during M6.")
    append_requirement(functional, "FR-23", "MUST", "Distribute mascotrender/0.1.0 with Conan 2 and install a relocatable CMake config exposing MascotRender::MascotRender.")
    append_requirement(functional, "FR-24", "MUST", "Provide an in-memory render API and a separate Conan test_package that compiles, links, and renders using only packaged artifacts.")

    nonfunctional = document.tables[9]
    update_requirement(nonfunctional, "NFR-02", "Determinism", "Byte-identical output is required within the same pinned Conan profiles, lockfile, fonts, renderer settings, and inputs. Other supported profiles require visual equivalence.")
    update_requirement(nonfunctional, "NFR-05", "Portability", "CMake and Conan must preserve Linux, macOS, and Windows packages. Linux and macOS are initial release gates; Windows is required by M5 unless explicitly waived.")
    update_requirement(nonfunctional, "NFR-07", "Security", "Engine 0.1 accepts trusted repository-controlled packs, contains all paths under the pack root, and rejects external URLs. Hostile third-party SVG isolation is a later security milestone.")

    # Technology baseline.
    technology = document.tables[10]
    tech_rows = [
        ("Language / Build", "C++20 and CMake", "Modern API, broad tooling, and a normal install/export package."),
        ("Package manager", "Conan 2", "Pinned dependencies, binary package IDs, host/build profiles, lockfiles, and external consumer testing."),
        ("CLI", "CLI11", "Small subcommand layer; only build and validate block engine 0.1."),
        ("Vector / Raster", "ThorVG 0.15.16 software backend", "Available through Conan Center and sufficient for SVG, vector primitives, transforms, text, and RGBA rendering."),
        ("Text", "ThorVG text + packaged static TTF for 0.1", "Exact English/Pidgin text, balanced wrapping, largest-valid whole-point fitting, and deterministic outlined glyph passes; shaping and fallback remain deferred."),
        ("WebP export", "libwebp 1.6.0", "Direct control of alpha and explicit deterministic encoding settings."),
        ("JSON", "nlohmann_json", "Simple schema/config parsing; parsing throughput is not an MVP bottleneck."),
        ("Image QA", "Owned basic validators", "Bounds, alpha, dimensions, and text checks without OpenCV; similarity analysis follows in M6."),
    ]
    # Reuse the existing eight data rows, then append the test row.
    for index, values in enumerate(tech_rows, start=1):
        set_row(technology, index, values)
    new_row = technology.add_row()
    for cell, value in zip(new_row.cells, ("Tests", "Catch2 + CTest + Conan test_package", "Unit/golden/integration coverage plus proof that installed artifacts are consumable."), strict=True):
        set_cell(cell, value)

    # Component and data design.
    replace_paragraph(
        document,
        "The loader parses CSV or JSON into a canonical StickerSpec. It validates required fields, normalizes locale and trigger values, verifies referenced assets, and computes a source fingerprint. Parsing errors must include file name, row number, column, bad value, and remediation guidance.",
        "Engine 0.1 parses schema-versioned JSON into a canonical StickerSpec. It "
        "validates required fields, normalizes locale and trigger values, verifies "
        "referenced assets, and reports the file plus JSON path, bad value, and "
        "remediation. A CSV importer is added in M6 and must produce this same model.",
    )
    replace_paragraph(
        document,
        "Asset imports must use a sanitized SVG subset; scripts, external URLs, embedded HTML, and remote fonts are rejected.",
        "Engine 0.1 accepts only trusted repository-controlled SVG parts, canonicalizes "
        "paths under the pack root, and rejects external URLs, embedded HTML, remote "
        "fonts, and disallowed references. A hardened untrusted-SVG boundary is later scope.",
    )
    replace_paragraph(
        document,
        "The engine, not the artwork generator, renders all sticker wording. This guarantees spelling and makes localization auditable. The layout service shapes glyphs, applies fallback fonts, measures the result, and searches for the largest valid font size within the configured text frame.",
        "The engine, not the artwork generator, renders all sticker wording. Engine "
        "0.1 loads only pack-declared font files and implements exact English/Pidgin "
        "measurement, balanced wrapping, largest-valid font-size search, and "
        "outline-aware safe-area fitting. HarfBuzz/FreeType "
        "shaping, locale fallback, bidirectional text, and Korean coverage are added "
        "after the base render pipeline is proven.",
    )
    replace_paragraph(
        document,
        "Locale-specific font fallback is configured in metadata. Fonts are packaged or referenced only when commercial distribution rights are confirmed.",
        "Every font is declared by content hash in the pack and may be packaged only after commercial distribution rights are confirmed. Locale-specific fallback is an M6+ extension.",
    )
    replace_paragraph(
        document,
        "The first implementation should expose an IRenderBackend interface and use Skia as the recommended production backend. The backend receives a CompiledScene and produces RGBA output. libwebp then performs final encoding. The separation permits a smaller alternative backend later without changing content or scene-compilation logic.",
        "Engine 0.1 uses an internal IRenderBackend implemented by the ThorVG software "
        "renderer pinned through Conan. It receives a CompiledScene and produces an "
        "owned RGBA buffer; libwebp performs final encoding. Backend types remain "
        "private so Skia or another renderer can be evaluated later without changing "
        "the public API or content model.",
    )

    replace_paragraph(document, "8.1 Authoring CSV", "8.1 Authoring JSON")
    replace_paragraph(
        document,
        "The existing CSV is extended rather than replaced. The first six columns remain understandable to non-engineers; advanced columns may be left blank and resolved from category defaults.",
        "Engine 0.1 uses schema-versioned JSON as the canonical authoring and test "
        "format. The M6 CSV importer is an adapter for content operators and must not "
        "create a second semantic model.",
    )
    set_code_cell(
        document.tables[18].cell(0, 0),
        '{\n  "schemaVersion": 1,\n  "stickers": [\n    {\n      "id": "stk_lol_001",\n      "text": "LOL",\n      "triggers": ["lol", "lmao", "laugh"],\n      "locale": "en",\n      "category": "laughter",\n      "emotion": "laughing",\n      "pose": "bent_laughing",\n      "effect": "tears",\n      "palette": "sunny",\n      "textStyle": "bold_bubble",\n      "seed": 101,\n      "priority": 100,\n      "enabled": true\n    }\n  ]\n}',
    )
    set_code_cell(
        document.tables[19].cell(0, 0),
        '{\n  "schemaVersion": 1,\n  "id": "wahalao-default",\n  "version": "0.1.0",\n  "canvas": {"width": 512, "height": 512, "safeMargin": 24},\n  "font": {"file": "fonts/sticker.ttf", "sha256": "..."},\n  "parts": {"head": "parts/head.svg", "body": "parts/body.svg"},\n  "anchors": {"leftHand": [142, 276], "rightHand": [370, 276]},\n  "expressions": {"laughing": "expressions/laughing.json"},\n  "poses": {"waving": "poses/waving.json"},\n  "effects": {"sparkles": "effects/sparkles.json"},\n  "textStyles": {"bold_bubble": "text/bold_bubble.json"},\n  "palettes": {"sunny": "palettes/sunny.json"}\n}',
    )

    # Algorithms and interfaces.
    replace_paragraph(document, "9.5 Incremental Build Cache", "9.5 Incremental Build Cache (M6)")
    set_code_cell(
        document.tables[23].cell(0, 0),
        "mascotrender build \\\n+  --input content/stickers.json \\\n+  --mascot assets/mascots/wahalao-default \\\n+  --output dist/v0.1.0 \\\n+  --strict\n\nmascotrender validate \\\n+  --input content/stickers.json \\\n+  --mascot assets/mascots/wahalao-default \\\n+  --strict\n\n# preview, inspect, package, and CSV import are M6+ workflow commands",
    )
    set_code_cell(
        document.tables[24].cell(0, 0),
        "namespace mascotrender {\n\n"
        "struct Error {\n"
        "    ErrorCode code;\n"
        "    std::string message;\n"
        "    std::filesystem::path source;\n"
        "    std::string location;\n"
        "};\n\n"
        "template<class T> class Result;\n\n"
        "struct RenderOptions {\n"
        "    std::uint32_t width{512};\n"
        "    std::uint32_t height{512};\n"
        "    float webpQuality{90.0F};\n"
        "    bool lossless{false};\n"
        "};\n\n"
        "struct EncodedImage {\n"
        "    std::uint32_t width{};\n"
        "    std::uint32_t height{};\n"
        "    std::string mediaType;\n"
        "    std::vector<std::byte> bytes;\n"
        "};\n\n"
        "class Pack;\n\n"
        "class MASCOTRENDER_API Engine {\n"
        "public:\n"
        "    Engine();\n"
        "    ~Engine();\n"
        "    Engine(Engine&&) noexcept;\n"
        "    Engine& operator=(Engine&&) noexcept;\n\n"
        "    Result<Pack> loadPack(const std::filesystem::path& path);\n"
        "    Result<EncodedImage> render(const Pack&, const StickerSpec&,\n"
        "                                const RenderOptions& = {});\n"
        "    Result<BuildSummary> build(const BuildRequest&);\n\n"
        "private:\n"
        "    class Impl;\n"
        "    std::unique_ptr<Impl> impl_;\n"
        "};\n\n"
        "} // namespace mascotrender",
    )

    insert_before(
        document,
        "7. Component Design",
        [
            ("6.4 Library Distribution Boundary", "Heading 2"),
            ("The engine is built as an ordinary CMake target and distributed as mascotrender/0.1.0 through Conan 2. The installed package exports MascotRender::MascotRender and contains public headers, library artifacts, licenses, and relocatable CMake config/version files.", "Normal"),
            ("Public headers use only standard-library and MascotRender-owned types; renderer, codec, JSON, and CLI dependencies remain private implementation details.", "List Bullet"),
            ("A separate Conan test_package must compile, link, run, load an example pack, and render using only packaged artifacts.", "List Bullet"),
            ("Conan package IDs capture OS, architecture, compiler, runtime, build type, and shared/static configuration. Cross-compiler C++ ABI compatibility is not promised.", "List Bullet"),
        ],
    )

    # Performance and determinism.
    performance = document.tables[28]
    perf_rows = [
        ("Engine clean batch", "50 representative stickers in <=60 seconds on the reference workstation", "M5/M6 CI benchmark"),
        ("Single sticker render", "<=500 ms excluding process startup on the reference workstation", "M3 benchmark"),
        ("Peak engine memory", "<=512 MB for the 50-sticker coherence build", "Process RSS"),
        ("Conan package", "Clean package plus test_package completes in <=10 minutes on each supported CI profile", "M5 package job"),
        ("Pilot per-keystroke lookup", "<=2 ms p95 for 10,000 triggers", "Browser microbenchmark in M7"),
        ("Pilot suggestion UI", "Tray updates within one animation frame when assets are cached", "React trace in M7"),
        ("Asset dimensions", "512 x 512 main, 256 x 256 thumbnail, transparent alpha", "Automated validator"),
        ("Asset size", "Target average <=150 KB for 512px static WebP", "Build report"),
        ("Visual defects", "Zero clipping, missing layers, opaque backgrounds, or unreadable text in approved release", "Automated QA plus human review"),
    ]
    for index, values in enumerate(perf_rows, start=1):
        set_row(performance, index, values)

    determinism_updates = {
        "Pin renderer, font, and dependency versions in the build image.": "Pin Conan host/build profiles, dependency versions/revisions, options, lockfiles, renderer settings, and font hashes.",
        "Use a fixed PRNG algorithm and seed derivation, not std::random_device.": "Use named fixed hash and PRNG algorithms; do not use std::hash, std::random_device, current time, or unordered iteration to determine output.",
        "Store source/output hashes and compare golden assets in CI.": "Require byte-identical output within the same pinned profile. Across different supported profiles require reviewed visual equivalence rather than identical encoded bytes.",
    }
    for old, new in determinism_updates.items():
        replace_paragraph(document, old, new)

    testing = document.tables[29]
    test_rows = [
        ("Unit tests", "JSON parsing, diagnostics, IDs, trigger normalization, transforms, text fitting, deterministic hash/PRNG, path containment, and validators."),
        ("Golden-image tests", "Ten representative renders compared within one pinned profile; cross-profile review uses explicit visual tolerances."),
        ("Property/fuzz tests", "Long strings, punctuation, Unicode, empty fields, extreme transforms, duplicate IDs, traversal, external URLs, and malformed JSON/SVG."),
        ("Integration tests", "End-to-end JSON to WebP/catalogue/dictionary/report and atomic failure behavior."),
        ("Package consumer test", "A separate Conan test_package locates MascotRender::MascotRender, links, runs, loads a pack, and renders."),
        ("Compatibility tests", "GCC, Clang/AppleClang, and MSVC package smoke tests across supported Conan profiles."),
        ("Performance tests", "Single render, ten/50-sticker batch, package time, and peak RSS; 1,000-sticker scale follows M6."),
        ("Security/content tests", "Trusted-pack boundary, path/URL rejection, size limits, license metadata, exact text, and human art review."),
    ]
    for index, values in enumerate(test_rows, start=1):
        set_row(testing, index, values)

    replace_paragraph(
        document,
        "The MVP should begin with approximately 200 high-frequency stickers rather than immediately publishing 1,000. The pilot must cover greetings, agreement, laughter, surprise, support, affection, status, work/school, food/sleep, frustration, celebration, and selected Pidgin expressions. The existing 204-row catalogue is a suitable planning baseline, subject to trigger cleanup and art prioritization.",
        "After engine 0.1.0, M6 produces a 50-sticker coherence set before scaling. "
        "Only after Product and Design approve that contact sheet does M7 produce "
        "approximately 200 high-frequency pilot stickers across the agreed categories. "
        "The existing 204-row catalogue is planning input, subject to trigger cleanup "
        "and art prioritization.",
    )
    replace_paragraph(
        document,
        "1.  Content change is submitted through version control with CSV/JSON and mascot-pack changes.",
        "1.  Content change is submitted through version control as canonical JSON and mascot-pack changes; the M6 CSV adapter produces the same model.",
    )

    # Delivery plan and risks.
    replace_paragraph(
        document,
        "The following is a planning baseline for a small cross-functional team. Durations are estimates and should be re-baselined after the renderer proof of concept and art-system workshop.",
        "M0-M5 are the installable engine MVP; M6-M7 are the product pilot. Estimates "
        "assume one senior C++ engineer with art work proceeding in parallel and must "
        "be re-baselined after M1 and the M6 coherence review. Detailed issue IDs and "
        "exit criteria are maintained in docs/MILESTONES.md.",
    )
    delivery = document.tables[30]
    delivery_rows = [
        ("M0 — Build/package bootstrap", "1-2 days", "C++20/CMake, Conan recipe, install/export, test_package, CI smoke", "External consumer installs, links, and calls the library."),
        ("M1 — First transparent sticker", "2-3 days", "ThorVG RGBA path, libwebp, in-memory render API, first golden", "Test consumer renders one deterministic transparent WebP."),
        ("M2 — Data-driven mascot pack", "3-4 days", "JSON schemas, pack loader, transforms, poses/expressions, diagnostics", "JSON content reproduces the golden and invalid inputs fail clearly."),
        ("M3 — Text/thumbnails/validation", "3-4 days", "Packaged font, fitting, main/thumb output, ten goldens", "Ten representative stickers pass exact-text and bounds gates."),
        ("M4 — Batch compiler/manifests", "2-3 days", "build/validate CLI, catalogue, dictionary, report, atomic output", "A clean versioned ten-sticker pack is deployable."),
        ("M5 — Engine 0.1 release", "2-3 days", "Compiler matrix, sanitizers, licenses, profiles/lockfiles, docs", "mascotrender/0.1.0 is published and adopted by a second consumer."),
        ("M6 — 50-sticker coherence gate", "1-2 weeks", "Art expansion, CSV adapter, cache/contact sheet as needed", "Product and Design approve 50 coherent stickers."),
        ("M7 — 200-sticker product pilot", "2-4 weeks", "Drogon/React/MLS integration, runtime QA, rollback, pilot content", "End-to-end pilot acceptance criteria pass."),
    ]
    for index, values in enumerate(delivery_rows, start=1):
        set_row(delivery, index, values)

    risks = document.tables[32]
    set_row(risks, 6, ("R-06", "Renderer or package dependency blocks portable Conan binaries", "Medium", "Pin ThorVG 0.15.16 for 0.1, keep backend types private, prove all supported profiles in M0/M5, and adopt a newer/custom recipe only against a measured requirement."))
    extra_risk = risks.add_row()
    for cell, value in zip(extra_risk.cells, ("R-11", "Engine scope expands into the product pilot", "High", "Treat M0-M5 and M6-M7 as separate release gates; deferred features cannot block engine 0.1 unless an accepted decision changes scope."), strict=True):
        set_cell(cell, value)

    # Acceptance criteria: retain product criteria under a separate heading.
    product_criteria = [
        "A clean build from the pinned environment generates the approved pilot pack without paid API calls.",
        "At least 200 approved stickers cover the agreed category matrix and use the original Wahalao mascot system.",
        "Every sticker has a stable ID, exact text, alt text, triggers, category, locale, dimensions, and content hash.",
        "No approved sticker has clipped text/art, missing alpha, unsafe background, unlicensed asset, or unresolved validation failure.",
        "The client makes no network request in response to each keystroke.",
        "The phrase matcher passes the trigger boundary suite, including no match for 'he' inside 'the' or 'weather'.",
        "Drogon endpoints support ETag, versioning, immutable asset caching, and rollback to the previous pack.",
        "The selected sticker ID is included in the encrypted outgoing message attachment contract and renders on a receiving client.",
        "Performance targets are benchmarked and any exceptions are documented and approved.",
        "Release documentation, license records, build report, and operational rollback steps are complete.",
    ]
    replace_paragraph(document, "17.1 MVP Acceptance Criteria", "17.1 Engine 0.1 Acceptance Criteria")
    engine_criteria = [
        "conan create . --build=missing succeeds from clean supported profiles and runs the separate test_package.",
        "An external C++20 consumer links MascotRender::MascotRender without including backend, codec, JSON, or CLI headers.",
        "The test consumer loads the example JSON mascot pack and renders transparent 512 x 512 and 256 x 256 WebP assets.",
        "Ten representative stickers render exact authored English/Pidgin text or fail with structured, source-located diagnostics.",
        "The same pinned Conan profiles, lockfile, inputs, font hashes, renderer settings, and seed produce byte-identical outputs.",
        "Missing references, duplicate IDs, path traversal, external URLs, text overflow, invalid dimensions, and missing alpha fail validation.",
        "build and validate produce an atomic versioned directory with catalogue.json, dictionary.json, build-report.json, and content hashes.",
        "Static Release packages pass Linux and macOS; the agreed compiler/profile matrix passes before publication.",
        "All dependency and content licenses required for distribution are recorded and shipped with the package or pack.",
        "No unresolved blocker-level unit, integration, golden, package-consumer, sanitizer, or license failure remains.",
    ]
    for old, new in zip(product_criteria, engine_criteria, strict=True):
        # The original document uses curly quotes in one criterion.
        if old.startswith("The phrase matcher passes"):
            old = "The phrase matcher passes the trigger boundary suite, including no match for “he” inside “the” or “weather”."
        replace_paragraph(document, old, new)

    insert_before(
        document,
        "17.2 Definition of Done for a Sticker",
        [("17.2 Product Pilot Acceptance Criteria", "Heading 2")]
        + [(item, "List Bullet") for item in product_criteria]
    )
    replace_paragraph(document, "17.2 Definition of Done for a Sticker", "17.3 Definition of Done for a Sticker")

    # Roadmap, decisions, and repository layout.
    roadmap = document.tables[34]
    roadmap_rows = [
        ("M6", "CSV importer, incremental cache, contact-sheet review, optional perceptual similarity, and 50-sticker art-system validation."),
        ("M6+", "HarfBuzz/FreeType shaping, fallback fonts, Korean corpus, Unicode normalization, bidi, and locale-specific safe areas."),
        ("M7", "Drogon/CDN delivery, React local matching, MLS attachment integration, rollback, and approved pilot telemetry."),
        ("Post-pilot", "Animated WebP/APNG/Lottie export from pose keyframes and procedural effects."),
        ("Post-pilot", "Additional mascot packs, seasonal themes, category lazy loading, and compact Trie/FST."),
        ("Later", "WASM preview, controlled personalization, and offline ML-assisted trigger expansion."),
        ("Later", "Technical-artist desktop editor built on the same scene graph and schemas."),
    ]
    for index, values in enumerate(roadmap_rows, start=1):
        set_row(roadmap, index, values)

    replace_paragraph(document, "19. Open Decisions", "19. Decisions and Remaining Open Items")
    decisions = document.tables[35]
    set_row(decisions, 0, ("ID", "Decision", "Status / Owner", "Needed By"))
    set_cell(decisions.rows[1].cells[3], "Before M2 exit")
    set_row(decisions, 2, ("OD-02", "First renderer is ThorVG software backend; libwebp performs encoding", "RESOLVED — C++ Lead", "Accepted in v0.2"))
    set_cell(decisions.rows[3].cells[3], "Before M3")
    set_cell(decisions.rows[4].cells[3], "Before M3")
    set_cell(decisions.rows[5].cells[3], "Before M7")
    set_row(decisions, 6, ("OD-06", "Engine 0.1 uses JSON dictionary only; binary Trie is post-pilot", "RESOLVED — Web + C++ Leads", "Accepted in v0.2"))
    set_cell(decisions.rows[7].cells[3], "Before M7 telemetry")
    set_cell(decisions.rows[8].cells[3], "Before M6")
    new_decision = decisions.add_row()
    for cell, value in zip(new_decision.cells, ("OD-09", "Supported Conan compiler/profile matrix and internal package remote", "C++ Lead + Platform", "Before M5"), strict=True):
        set_cell(cell, value)

    set_code_cell(
        document.tables[36].cell(0, 0),
        "MascotRender/\n"
        "  CMakeLists.txt\n"
        "  conanfile.py\n"
        "  conandata.yml\n"
        "  cmake/MascotRenderConfig.cmake.in\n"
        "  include/mascotrender/\n"
        "  src/{spec,pack,scene,text,render,export,validate,package}/\n"
        "  apps/mascotrender/\n"
        "  tests/{unit,integration,golden}/\n"
        "  test_package/\n"
        "  examples/basic/\n"
        "  assets/example-pack/\n"
        "  content/example-stickers.json\n"
        "  profiles/\n"
        "  docs/{MILESTONES.md,DECISIONS.md}/\n"
        "  tools/",
    )
    document.tables[38].cell(0, 0).text = (
        "Prove installation, then rendering, then coherence, then volume\n"
        "The first checkpoint is a Conan test consumer that renders one transparent "
        "WebP. The engine checkpoint is ten representative stickers and a publishable "
        "mascotrender/0.1.0 package. The next executive checkpoint is a 50-sticker "
        "contact sheet. Only after Design, Product, and Engineering approve coherence "
        "should the team scale toward the approximately 200-sticker pilot or 1,000 assets."
    )

    # Appendix D makes the package contract self-contained.
    document.add_heading("Appendix D — Conan and CMake Package Contract", level=1)
    document.add_paragraph(
        "Canonical Conan reference: mascotrender/0.1.0. Canonical CMake package and "
        "target: MascotRender and MascotRender::MascotRender. The recipe owns "
        "dependency resolution; the CMake project contains no Conan-specific logic "
        "and remains directly installable.",
    )
    package_table = document.add_table(rows=1, cols=2)
    package_table.style = "Table Grid"
    set_row(package_table, 0, ("Contract", "Engine 0.1 Requirement"))
    for left, right in [
        ("Recipe", "Conan 2 conanfile.py with library package_type, shared/fPIC options, settings, requirements, layout, build, package, and package_info."),
        ("CMake", "Install headers, library artifacts, licenses, MascotRenderConfig.cmake, version file, and exported targets using relative relocatable paths."),
        ("Consumer", "find_package(MascotRender CONFIG REQUIRED) followed by target_link_libraries(app PRIVATE MascotRender::MascotRender)."),
        ("Proof", "test_package is built outside the source target graph and renders using only packaged artifacts."),
        ("Determinism", "Release artifacts record Conan profile names, lockfile, recipe/dependency revisions, options, renderer settings, and font/content hashes."),
        ("ABI", "Conan package ID separates compiler, runtime, OS, architecture, build type, and shared/static variants; no cross-compiler ABI promise."),
    ]:
        row = package_table.add_row()
        set_cell(row.cells[0], left)
        set_cell(row.cells[1], right)

    document.add_paragraph()
    consumer = document.add_paragraph()
    consumer.style = "Normal"
    run = consumer.add_run(
        "find_package(MascotRender CONFIG REQUIRED)\n"
        "target_link_libraries(my_app PRIVATE MascotRender::MascotRender)"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    document.add_heading("Appendix E — 0.1 Release Candidate Status", level=1)
    document.add_paragraph(
        "As of 13 July 2026, M0-M4 are implemented locally. The engine renders "
        "five deterministic mascot identities and 50 stickers through the C++20 "
        "library, with 512px assets, thumbnails, catalogue, dictionary, build "
        "report, balanced text, outlined glyphs, and a decoded-pixel golden. The "
        "AppleClang Release suite passes 18 tests, and external static/shared Conan "
        "consumers pass. The 50-sticker render baseline is 4.84 seconds on the "
        "recorded macOS arm64 workstation."
    )
    document.add_paragraph(
        "The remaining M5 release gates are passing hosted GCC, Clang, MSVC, ASan, "
        "and UBSan jobs; owner-supplied project license text; and selection plus "
        "authentication of a writable Conan remote. Product and Design approval of "
        "the complete generated art set remains the independent M6 coherence gate."
    )

    # Header and core properties.
    for section in document.sections:
        for item in section.header.paragraphs:
            if "MascotRender SDD" in item.text:
                item.text = "MascotRender SDD  |  Implementation Baseline v0.2"
    document.core_properties.title = "MascotRender Software Design Document v0.2"
    document.core_properties.subject = (
        "Implementation baseline for the installable C++20 mascot rendering engine"
    )
    document.core_properties.comments = "Implementation baseline v0.2"

    document.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    main()
